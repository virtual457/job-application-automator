"""
Simple Resume Generator

Reads content from config/current_application.yaml
Applies it to templates/Chandan_Resume_Format.docx
Saves to output/Generated_Resume.docx

Usage:
    python simple_generator.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import yaml
from pathlib import Path
import os
import time
import re


def add_text_with_bold_markers(paragraph, text, font_size=10, base_bold=False):
    """
    Add text with **bold** markers.
    Example: "Built **AWS Lambda** pipeline" [*] "Built AWS Lambda pipeline" (with AWS Lambda bold)
    
    Args:
        paragraph: The Word paragraph to add text to
        text: Text with **markers** for bold sections
        font_size: Font size for all text
        base_bold: If True, non-marked text is also bold (marked text becomes non-bold)
    """
    parts = text.split('**')
    
    for i, part in enumerate(parts):
        if not part:  # Skip empty parts
            continue
            
        run = paragraph.add_run(part)
        run.font.size = Pt(font_size)
        
        # Even indices (0, 2, 4...) = normal text
        # Odd indices (1, 3, 5...) = text between ** markers = bold
        if i % 2 == 1:
            run.bold = True
        else:
            run.bold = base_bold


def add_hyperlink(paragraph, text, url, font_size=None, bold=False, italic=False):
    """
    Add a hyperlink to a paragraph.
    Returns the hyperlink run so you can format it.
    """
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True))
    
    # Create a new run for the hyperlink
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add bold if specified
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    
    # Add italic if specified
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)
    
    # Add underline and color (blue) for hyperlink style
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')  # Blue color
    rPr.append(color)
    
    # Add font size if specified
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))  # Word uses half-points
        rPr.append(sz)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._element.append(hyperlink)
    
    return hyperlink


def find_paragraph_by_text(doc, search_text):
    """Find a paragraph that contains the search text"""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            return i
    return None


def close_all_word_applications():
    """Close all Word applications using taskkill"""
    print("[*] Closing all Word applications...")
    try:
        os.system("taskkill /F /IM WINWORD.EXE >nul 2>&1")
        time.sleep(1.5)
        print("   [OK] Word closed\n")
        return True
    except Exception as e:
        print(f"   [!]  Could not close Word: {e}\n")
        return False


def update_header(doc, header_data):
    """Update the first 3 paragraphs (name, title, contact with hyperlinks)"""
    print("[*] Updating header...")
    
    if len(doc.paragraphs) >= 3:
        # Line 1: Name
        doc.paragraphs[0].text = header_data['name']
        
        # Line 2: Title  
        doc.paragraphs[1].text = header_data['title']
        
        # Line 3: Contact info with hyperlinks
        contact_para = doc.paragraphs[2]
        contact_para.clear()
        
        # Phone number (plain text)
        run = contact_para.add_run("+1 (857) 421-7469; ")
        run.font.size = Pt(10)
        
        # Email (hyperlink)
        add_hyperlink(contact_para, "chandan.keelara@gmail.com", "mailto:chandan.keelara@gmail.com", font_size=10)
        contact_para.add_run("; ")
        
        # LinkedIn (hyperlink)
        add_hyperlink(contact_para, "LinkedIn", "https://www.linkedin.com/in/chandan-gowda-k-s-765194186/", font_size=10)
        contact_para.add_run("; ")
        
        # Portfolio (hyperlink)
        add_hyperlink(contact_para, "Portfolio", "https://virtual457.github.io/", font_size=10)
        contact_para.add_run("; ")
        
        # GitHub (hyperlink)
        add_hyperlink(contact_para, "GitHub", "https://github.com/virtual457", font_size=10)
        contact_para.add_run(";")
        
        print("   [OK] Header updated (with hyperlinks)")
    else:
        print("   [!]  Warning: Not enough paragraphs for header")


def update_summary(doc, summary_text):
    """Find and update the summary paragraph with bold marker support"""
    print("[*] Updating summary...")
    
    # Find the paragraph that starts with "Software Engineer and MS CS student"
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("Software Engineer and MS CS student"):
            # Clear existing runs
            para.clear()
            # Add text with bold marker support
            add_text_with_bold_markers(para, summary_text, font_size=10)
            print("   [OK] Summary updated (with bold markers)")
            return
    
    print("   [!]  Warning: Could not find summary paragraph")


def update_skills(doc, skills_list):
    """Update the TECHNICAL SKILLS section with dynamic tab alignment"""
    print("[*] Updating technical skills...")
    
    # Find "TECHNICAL SKILLS" heading
    skills_idx = find_paragraph_by_text(doc, "TECHNICAL SKILLS")
    
    if skills_idx is None:
        print("   [!]  Warning: Could not find TECHNICAL SKILLS section")
        return
    
    # Fixed tab position at 35 characters for consistent alignment
    # Average char width in 10pt Calibri bold - using 0.0629 inches per char
    tab_position = 2.2  # 2.2 inches
    
    print(f"   [*] Tab position: {tab_position:.2f} inches (fixed at 35 chars)")
    
    # Update the next 7 paragraphs with skills
    for i, skill in enumerate(skills_list):
        para_idx = skills_idx + 1 + i
        if para_idx < len(doc.paragraphs):
            para = doc.paragraphs[para_idx]
            para.clear()
            
            # Set tab stop at calculated position
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(tab_position), WD_TAB_ALIGNMENT.LEFT)
            
            # Add category (bold)
            category_run = para.add_run(skill['category'])
            category_run.bold = True
            category_run.font.size = Pt(10)
            
            # Add single tab (will align at calculated position)
            para.add_run("\t")
            
            # Add items (not bold)
            items_run = para.add_run(skill['items'])
            items_run.bold = False
            items_run.font.size = Pt(10)
    
    print(f"   [OK] Updated {len(skills_list)} skill categories (dynamic alignment)")


def update_experience(doc, experience_list):
    """Update the WORK EXPERIENCE section with dynamic bullets"""
    print("[*] Updating work experience...")
    
    # Find "WORK EXPERIENCE" heading
    exp_idx = find_paragraph_by_text(doc, "WORK EXPERIENCE")
    
    if exp_idx is None:
        print("   [!]  Warning: Could not find WORK EXPERIENCE section")
        return
    
    current_para = exp_idx + 1
    
    for company_exp in experience_list:
        # Skip the company header line (it stays unchanged in template)
        # The template has: "Company | Title | Location    Duration"
        current_para += 1
        
        # Update bullets for this company
        num_bullets = len(company_exp['bullets'])
        
        for bullet_text in company_exp['bullets']:
            if current_para >= len(doc.paragraphs):
                print(f"   [!]  Warning: Not enough paragraphs for {company_exp['company']} bullets")
                break
            
            # Update bullet paragraph with bold marker support
            para = doc.paragraphs[current_para]
            para.clear()
            add_text_with_bold_markers(para, bullet_text, font_size=10)
            current_para += 1
        
        print(f"   [OK] Updated {num_bullets} bullets for {company_exp['company']}")
    
    print(f"   [OK] Work experience section complete")


def update_projects(doc, projects_list):
    """Update the PROJECTS section with GitHub hyperlinks"""
    print("[*] Updating projects...")
    
    # GitHub URLs for projects
    github_urls = {
        "LLM Multi-Agent Resume Optimizer": "https://github.com/virtual457/llm-multi-agent-resume-optimizer",
        "LLM Multi-Agent Resume Optimizer (LMARO)": "https://github.com/virtual457/llm-multi-agent-resume-optimizer",
        "LMARO": "https://github.com/virtual457/llm-multi-agent-resume-optimizer",
        "Dino Game Deep RL Agent": "https://github.com/virtual457/dino-game-AI",
        "Orion PaaS": "https://github.com/virtual457/Orion-platform",
        "Orion Platform": "https://github.com/virtual457/Orion-platform",
        "Calendly - Calendar Management System": "https://github.com/virtual457/Calendly",
        "Calendly": "https://github.com/virtual457/Calendly",
        "Maritime Logistics Platform": "https://github.com/virtual457/Port-Management-System",
        "Port Management System": "https://github.com/virtual457/Port-Management-System",
        "Large Scale Data Analysis": "https://github.com/virtual457/Data-analysis-on-pubg",
        "Data Analysis on PUBG": "https://github.com/virtual457/Data-analysis-on-pubg",
        "Face Recognition & Validation System": "https://github.com/virtual457/Recognition-and-Validation-of-Faces-using-Machine-Learning-and-Image-Processing",
        "Online Examination System": "https://github.com/virtual457/Online-examination-using-mongodb",
        "Kambaz Learning Management System": "https://github.com/virtual457/kambaz-next-js",
        "Healthcare Management System": "https://github.com/virtual457/stuffycare",
        "StuffyCare": "https://github.com/virtual457/stuffycare"
    }
    
    # Find "PROJECTS" heading
    projects_idx = find_paragraph_by_text(doc, "PROJECTS")
    
    if projects_idx is None:
        print("   [!]  Warning: Could not find PROJECTS section")
        return
    
    current_para = projects_idx + 1
    
    for project in projects_list:
        if current_para >= len(doc.paragraphs):
            print(f"   [!]  Warning: Not enough paragraphs for project: {project['title']}")
            break
        
        # Update title line with hyperlinked GitHub
        para = doc.paragraphs[current_para]
        para.clear()
        
        # Add project title as hyperlink (BOLD)
        # Use github field from YAML if it's a URL, otherwise fall back to dictionary
        github_url = None
        if 'github' in project and project['github'].startswith('http'):
            github_url = project['github']
        elif project['title'] in github_urls:
            github_url = github_urls[project['title']]
        
        if github_url:
            add_hyperlink(para, project['title'], github_url, font_size=10, bold=True)
        else:
            title_run = para.add_run(project['title'])
            title_run.bold = True
            title_run.font.size = Pt(10)
        
        para.add_run(" | ")
        
        # Add tech stack (ITALIC ONLY)
        tech_run = para.add_run(project['tech'])
        tech_run.bold = False
        tech_run.italic = True
        tech_run.font.size = Pt(10)
        
        para.add_run(" | ")
        
        # Add GitHub text (also hyperlink) - use same URL
        if github_url:
            add_hyperlink(para, "GitHub", github_url, font_size=10)
        else:
            github_run = para.add_run("GitHub")
            github_run.font.size = Pt(10)
        
        current_para += 1
        
        # Update bullet 1 (template already has bullet)
        if current_para < len(doc.paragraphs):
            para = doc.paragraphs[current_para]
            para.clear()
            add_text_with_bold_markers(para, project['bullet1'], font_size=10)
            current_para += 1
        
        # Update bullet 2 (template already has bullet)
        if current_para < len(doc.paragraphs):
            para = doc.paragraphs[current_para]
            para.clear()
            add_text_with_bold_markers(para, project['bullet2'], font_size=10)
            current_para += 1
    
    print(f"   [OK] Updated {len(projects_list)} projects (with GitHub hyperlinks)")


def open_in_word(file_path):
    """Open the generated resume in Word"""
    print("\n[*] Opening resume in Word...")
    try:
        os.startfile(str(file_path))
        print("   [OK] Resume opened in Word")
        return True
    except Exception as e:
        print(f"   [!]  Could not open Word: {e}")
        return False


def extract_company_name(content):
    """
    Extract target company name from YAML content.
    Looks for 'company_name' field first, then tries to extract from summary.
    """
    # Check if company_name field exists
    if 'company_name' in content:
        return content['company_name']

    # Try to extract from summary (e.g., "Excited to contribute to Intuit's mission")
    summary = content.get('summary', '')

    # Pattern 1: "contribute to [Company]'s"
    match = re.search(r"contribute to ([A-Za-z0-9\s&]+)'s", summary)
    if match:
        return match.group(1).strip()

    # Pattern 2: "join [Company]'s team"
    match = re.search(r"join ([A-Za-z0-9\s&]+)'s team", summary)
    if match:
        return match.group(1).strip()

    # Pattern 3: "at [Company]"
    match = re.search(r"at ([A-Za-z0-9\s&]+)$", summary)
    if match:
        return match.group(1).strip()

    # Default fallback
    return "Company"


def sanitize_filename(name):
    """
    Convert company name to filename-friendly format.
    Example: "London Stock Exchange Group (LSEG)" -> "london_stock_exchange_group_lseg"
    """
    # Remove special characters and convert to lowercase
    name = re.sub(r'[^\w\s-]', '', name.lower())
    # Replace spaces with underscores
    name = re.sub(r'[\s-]+', '_', name)
    # Remove multiple underscores
    name = re.sub(r'_+', '_', name)
    return name.strip('_')


def run_validator():
    """
    Run the YAML validator (validate_yaml.py).
    Returns True if validation passes, False otherwise.
    """
    import subprocess

    print("\n[VALIDATE] Running YAML validation...")

    try:
        # Run the validator script
        result = subprocess.run(
            ["python", "validate_yaml.py"],
            capture_output=True,
            text=True,
            timeout=30
        )

        # Print validator output
        if result.stdout:
            print(result.stdout)

        # Check if validation passed
        if result.returncode == 0:
            print("   [OK] Validation PASSED")
            return True
        else:
            print("   [X] Validation FAILED")
            if result.stderr:
                print(f"   [!] Error: {result.stderr}")
            return False

    except subprocess.TimeoutExpired:
        print("   [!] Validation timed out")
        return False
    except Exception as e:
        print(f"   [!] Could not run validator: {e}")
        return False


def convert_to_pdf(docx_path, company_name):
    """
    Convert DOCX to PDF and save to D:\\Resumes folder.
    Filename format: chandan_resume_companyname.pdf
    """
    try:
        from docx2pdf import convert
    except ImportError:
        print("\n[!]  Warning: docx2pdf not installed")
        print("   [*] Install with: pip install docx2pdf")
        print("   [*] Skipping PDF generation")
        return False

    print("\n[PDF] Converting to PDF...")

    # Create D:\Resumes directory if it doesn't exist
    pdf_dir = Path("D:/Resumes")
    pdf_dir.mkdir(parents=True, exist_ok=True)

    # Create filename
    sanitized_company = sanitize_filename(company_name)
    pdf_filename = f"chandan_resume_{sanitized_company}.pdf"
    pdf_path = pdf_dir / pdf_filename

    try:
        # Convert DOCX to PDF
        convert(str(docx_path), str(pdf_path))
        print(f"   [OK] PDF saved to: {pdf_path}")
        return True
    except Exception as e:
        print(f"   [!]  PDF conversion failed: {e}")
        print("   [*] Make sure Microsoft Word is installed (required for docx2pdf)")
        return False


def generate_resume():
    """Main function to generate resume"""
    
    print("\n" + "="*60)
    print("SIMPLE RESUME GENERATOR")
    print("="*60 + "\n")
    
    # Close all Word applications FIRST
    close_all_word_applications()
    
    # Paths
    yaml_path = Path("config/current_application.yaml")
    template_path = Path("templates/Chandan_Resume_Format.docx")
    output_path = Path("output/Generated_Resume.docx")
    
    # Check if files exist
    if not yaml_path.exists():
        print(f"[X] Error: YAML file not found at {yaml_path}")
        return
    
    if not template_path.exists():
        print(f"[X] Error: Template not found at {template_path}")
        print("[*] Make sure Chandan_Resume_Format.docx is in templates/ folder")
        return
    
    # Create output directory
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Load content from YAML
    print("[*] Loading content from YAML...")
    with open(yaml_path, 'r') as f:
        content = yaml.safe_load(f)
    print("   [OK] Content loaded\n")
    
    # Load Word template
    print("[*] Loading Word template...")
    doc = Document(template_path)
    print(f"   [OK] Template loaded ({len(doc.paragraphs)} paragraphs)\n")
    
    # Update each section
    update_header(doc, content['header'])
    update_summary(doc, content['summary'])
    update_skills(doc, content['skills'])
    
    # Update work experience if present in YAML
    if 'experience' in content:
        update_experience(doc, content['experience'])
    
    update_projects(doc, content['projects'])
    
    # Save output
    print("\n[SAVE] Saving generated resume...")
    doc.save(str(output_path))
    print(f"   [OK] Saved to: {output_path}")

    # Run validator before PDF generation
    validation_passed = run_validator()

    # Extract company name and convert to PDF only if validation passed
    pdf_success = False
    company_name = extract_company_name(content)
    print(f"\n[*] Target company: {company_name}")

    if validation_passed:
        pdf_success = convert_to_pdf(output_path, company_name)
    else:
        print("\n[!] Skipping PDF generation due to validation errors")
        print("   [*] Fix validation errors and rerun to generate PDF")

    # Open in Word
    open_in_word(output_path)

    print("\n" + "="*60)
    print("[OK] RESUME GENERATION COMPLETE!")
    print("="*60)
    print(f"\n[*] DOCX Output: {output_path.absolute()}")
    if pdf_success:
        sanitized_company = sanitize_filename(company_name)
        pdf_path = Path(f"D:/Resumes/chandan_resume_{sanitized_company}.pdf")
        print(f"[*] PDF Output: {pdf_path.absolute()}")
    print("\n[*] Next steps:")
    print("   1. Review the formatting in Word")
    print("   2. Make manual adjustments if needed")
    if not validation_passed:
        print("   3. Fix validation errors in YAML")
        print("   4. Rerun to generate PDF")
    elif not pdf_success:
        print("   3. Install docx2pdf: pip install docx2pdf")
        print("   4. Rerun to generate PDF automatically")


if __name__ == "__main__":
    try:
        generate_resume()
    except Exception as e:
        print(f"\n[X] Error: {e}")
        import traceback
        traceback.print_exc()
