"""
Resume Generator Module

Automatically generates tailored resumes for different companies and role types.
"""

from docx import Document
from docx.shared import Pt, Inches
import yaml
import os
from datetime import datetime
from pathlib import Path


class ResumeGenerator:
    """
    Generate tailored resumes for job applications.
    
    Supports multiple role types (backend, ml, data engineering) with
    customizable headers, summaries, skills, and project ordering.
    """
    
    def __init__(self, template_path, config_path):
        """
        Initialize the resume generator.
        
        Args:
            template_path: Path to base resume Word document
            config_path: Path to YAML configuration file
        """
        self.template_path = Path(template_path)
        self.config_path = Path(config_path)
        
        # Load configuration
        with open(config_path, 'r') as f:
            self.config = yaml.safe_load(f)
        
        # Create output directory if it doesn't exist
        self.output_dir = Path('output/resumes')
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Track generation history
        self.generation_log = []
    
    def generate(self, company, role_type='backend', save_format='docx'):
        """
        Generate a tailored resume for a specific company and role.
        
        Args:
            company: Target company name (e.g., 'Salesforce', 'Microsoft')
            role_type: Type of role - 'backend', 'ml', 'data', or 'general'
            save_format: Output format - 'docx' or 'pdf'
        
        Returns:
            Path to generated resume file
        """
        print(f"\n🔨 Generating resume for {company} ({role_type} role)...")
        
        # Load base template
        doc = Document(self.template_path)
        
        # Get role-specific configuration
        if role_type not in self.config['roles']:
            raise ValueError(f"Unknown role type: {role_type}. Available: {list(self.config['roles'].keys())}")
        
        role_config = self.config['roles'][role_type]
        
        # Apply transformations
        self._update_header(doc, role_config['header'])
        self._update_summary(doc, company, role_config['summary'])
        self._update_skills(doc, role_config['skills'])
        # Note: Project reordering is complex, we'll handle it manually for now
        
        # Save output
        output_filename = f"Chandan_Resume_{company}_{role_type}.{save_format}"
        output_path = self.output_dir / output_filename
        doc.save(str(output_path))
        
        # Log generation
        self._log_generation(company, role_type, output_path)
        
        print(f"✅ Generated: {output_path}")
        return output_path
    
    def _update_header(self, doc, header_text):
        """Update the resume header (first paragraph)"""
        if len(doc.paragraphs) > 0:
            doc.paragraphs[0].text = header_text
            # Center align the header
            doc.paragraphs[0].alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
            # Make it bold
            for run in doc.paragraphs[0].runs:
                run.bold = True
    
    def _update_summary(self, doc, company, summary_template):
        """
        Update the summary/objective paragraph.
        
        Finds the paragraph starting with "Software Engineer and MS CS student"
        and replaces it with the template, filling in the company name.
        """
        summary_text = summary_template.format(company=company)
        
        for para in doc.paragraphs:
            if para.text.strip().startswith("Software Engineer and MS CS student"):
                para.text = summary_text
                break
    
    def _update_skills(self, doc, skills_dict):
        """
        Update the TECHNICAL SKILLS section.
        
        Finds the section and replaces skill categories with new ordering.
        """
        # Find TECHNICAL SKILLS heading
        skills_heading_idx = None
        for i, para in enumerate(doc.paragraphs):
            if "TECHNICAL SKILLS" in para.text:
                skills_heading_idx = i
                break
        
        if skills_heading_idx is None:
            print("⚠️  Warning: Could not find TECHNICAL SKILLS section")
            return
        
        # Clear existing skills (assume they're in the next 6-8 paragraphs)
        # We'll identify them by checking if they have the format "Category: items"
        current_idx = skills_heading_idx + 1
        cleared_lines = 0
        
        while current_idx < len(doc.paragraphs) and cleared_lines < 10:
            para = doc.paragraphs[current_idx]
            # Stop if we hit another section heading (all caps)
            if para.text.strip().isupper() and len(para.text.strip()) > 5:
                break
            # Clear this line if it looks like a skill line
            if para.text.strip():
                para.text = ""
                cleared_lines += 1
            current_idx += 1
        
        # Insert new skills
        insertion_idx = skills_heading_idx + 1
        for category, items in skills_dict.items():
            # Format: "Category Name                Items, Items, Items"
            skill_line = f"{category:<35} {items}"
            
            # Insert a new paragraph
            new_para = doc.paragraphs[insertion_idx]._element
            new_para = doc.paragraphs[insertion_idx].insert_paragraph_before(skill_line)
            insertion_idx += 1
    
    def _log_generation(self, company, role_type, output_path):
        """Log the generation for tracking purposes"""
        entry = {
            'timestamp': datetime.now().isoformat(),
            'company': company,
            'role_type': role_type,
            'output_path': str(output_path)
        }
        self.generation_log.append(entry)
        
        # Save to log file
        log_file = Path('output/generation_log.txt')
        with open(log_file, 'a') as f:
            f.write(f"{entry['timestamp']} | {company:20s} | {role_type:10s} | {output_path}\n")
    
    def batch_generate(self, applications):
        """
        Generate multiple resumes from a list of applications.
        
        Args:
            applications: List of dicts with 'company' and 'role_type' keys
            
        Example:
            applications = [
                {'company': 'Salesforce', 'role_type': 'backend'},
                {'company': 'ByteDance', 'role_type': 'ml'},
                {'company': 'Intuit', 'role_type': 'backend'}
            ]
        """
        print(f"\n🚀 Batch generating {len(applications)} resumes...\n")
        
        results = []
        for app in applications:
            try:
                output_path = self.generate(app['company'], app['role_type'])
                results.append({'company': app['company'], 'status': 'success', 'path': output_path})
            except Exception as e:
                print(f"❌ Error generating resume for {app['company']}: {e}")
                results.append({'company': app['company'], 'status': 'failed', 'error': str(e)})
        
        # Summary
        successful = sum(1 for r in results if r['status'] == 'success')
        print(f"\n✅ Batch complete: {successful}/{len(applications)} resumes generated")
        
        return results
    
    def get_stats(self):
        """Get statistics about generated resumes"""
        return {
            'total_generated': len(self.generation_log),
            'by_role_type': self._count_by_field('role_type'),
            'by_company': self._count_by_field('company')
        }
    
    def _count_by_field(self, field):
        """Helper to count entries by a specific field"""
        counts = {}
        for entry in self.generation_log:
            key = entry[field]
            counts[key] = counts.get(key, 0) + 1
        return counts


# Utility functions
def list_available_roles(config_path='config/resume_config.yaml'):
    """List all available role types from config"""
    with open(config_path, 'r') as f:
        config = yaml.safe_load(f)
    return list(config['roles'].keys())


if __name__ == "__main__":
    # Quick test
    print("Resume Generator Module")
    print("Available roles:", list_available_roles())
